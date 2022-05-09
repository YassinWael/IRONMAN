import os

import random
import shutil
from urllib.request import Request, urlopen
import docx
import pyttsx3
import speech_recognition
from datetime import datetime
import pywhatkit as kt
import pyautogui
import time as t
import winshell
from PyDictionary import PyDictionary
import webbrowser
import clipboard
from bs4 import BeautifulSoup
from docx.shared import Pt
from plyer import notification
if __name__ == '__main__':
    notification.notify(title='J.A.R.V.I.S ',
                        message='jarvis is up and running',
                        app_icon="C:\\Users\\Dr.Wael\\PycharmProjects\\The talking project\\icon.ico",
                        timeout=10)

today_date = datetime.now()
try:
    import vlc
except ValueError:
    print('failed to import vlc')
i = 0
engine = pyttsx3.init()
voices = engine.getProperty('voices')
engine.setProperty('voice', voices[0].id)
engine.setProperty('rate', 170)


def speak(s):  # the function that speaks

    engine.say(s)
    engine.runAndWait()





def get_audio(printorno=True):  # function that recognizes speech
    r = speech_recognition.Recognizer()
    with speech_recognition.Microphone() as source:
        r.adjust_for_ambient_noise(source, duration=6)
        if printorno is True:
            print("talk now please")
            speak('talk')

        try:

            audio = r.listen(source, timeout=5)
        except Exception:
            pass

        said = ""
        try:

            said = r.recognize_google(audio)

            print(said)
        except Exception as Error:
            pass

    return said


def neffex_new_song_download():
    today_date = datetime.now()
    day = today_date.strftime("%A")
    if day == 'Wednesday':
        speak('neffex just added a new video!')
        req = Request("https://soundcloud.com/neffexmusic")
        html_page = urlopen(req)
        soup = BeautifulSoup(html_page, "lxml")
        links = []
        for link in soup.findAll('a'):  # the link we need "https://soundcloud.com/neffexmusic/play-dead"
            links.append(link.get('href'))
        latest_song = f'https://soundcloud.com/{links[3]}'
        print(f'the latest song is {latest_song}')
        song_name = latest_song.split("/")
        speak('should i download it?')
        text = get_audio(printorno=True)
        if 'yes' in text:
            clipboard.copy(latest_song)
            webbrowser.open('https://scloudtomp3downloader.com/?lang=en')
            t.sleep(50)
            pyautogui.leftClick(x=424, y=281)
            pyautogui.hotkey('ctrl', 'v')
            t.sleep(2)
            pyautogui.leftClick(x=1011, y=287)
            t.sleep(6)
            pyautogui.leftClick(x=604, y=200)
            t.sleep(4)
            pyautogui.leftClick(x=826, y=480)
            t.sleep(2.5)
            pyautogui.leftClick(x=740, y=425)
            print(day)
            speak('downloaded it will try to rename')
            downloads = os.listdir("C:\\Users\\Dr.Wael\\Downloads")
            for d in downloads:
                if '128' in d:
                    t.sleep(30)
                    print(d)
                    os.rename(f"C:\\Users\\Dr.Wael\\Downloads\\{d}",
                              f"C:\\Users\\Dr.Wael\\PycharmProjects\\The talking project\\musicfiles\\{song_name[5]}.mp3")
                    speak('we did it')


neffex_new_song_download()


def lock_keyboard():
    os.startfile(r"C:\\Users\\Dr.Wael\\Downloads\\Keyboard-Locker\\Keyboard Locker\\KeyboardLocker.exe")
    t.sleep(1)
    pyautogui.hotkey('ctrl', 'alt', 'l')
    speak('keyboard locked')


def unlock_keyboard():
    speak('unlocking')
    pyautogui.write('unlock')
    speak('unlocked')


def file_delete():
    value_of_files = 0
    path = r"C:\\Users\\Dr.Wael\\Desktop\\python test example\\\\"
    destination = "C:\\\\Users\\\\Dr.Wael\\\\Desktop\\\\file to always recycle"
    files_in_destination = os.listdir(path)
    print(files_in_destination)
    for file in files_in_destination:
        if file.endswith('.jpg'):
            value_of_files = value_of_files + 1
    speak(f'the amount of photo files in this folder is {value_of_files}')

    speak('what do you want to do with those files,sir?')
    text = get_audio(printorno=True)
    if 'open' in text:
        for file_name in files_in_destination:
            speak(f'opening {file_name}')
            print(f"C:\\\\Users\\\\Dr.Wael\\\\Desktop\\\\python test example\\\\{file_name}")
            os.startfile(f"C:\\\\Users\\\\Dr.Wael\\\\Desktop\\\\python test example\\\\{file_name}")
            t.sleep(4)
    if 'delete' in text:
        speak('deleting')
        for file_name in files_in_destination:
            print(file_name)
            shutil.move(os.path.join(path, file_name), destination)
            speak('finished moving')
        t.sleep(10)
        os.startfile(r"C:\\Users\\Dr.Wael\\Desktop\\file to always recycle")
        t.sleep(5)
        pyautogui.leftClick(x=977, y=368)
        t.sleep(1)
        pyautogui.hotkey('ctrl', 'a')
        t.sleep(2)
        pyautogui.press('del')
        speak('finished deleting')


def call_phone():
    speak('calling the phone')
    webbrowser.open('https://myaccount.google.com/find-your-phone?continue=https://myaccount.google.com/?utm_source'
                    '%3Daccount-marketing-page%26utm_medium%3Dgo-to-account-button%26pli%3D1')
    t.sleep(13)
    pyautogui.leftClick(601, 446)
    t.sleep(2)
    pyautogui.leftClick(225, 412)
    speak('called')


def get_song_by_lyrics():
    speak('please say a part of the lyrics')
    lyrics = get_audio(printorno=True)
    webbrowser.open('https://songsear.ch')
    t.sleep(5)
    pyautogui.leftClick(x=568, y=270)
    t.sleep(1)
    pyautogui.write(lyrics)
    pyautogui.press('enter')
    pyautogui.leftClick(x=608, y=52)
    pyautogui.hotkey('ctrl', 'c')
    website = clipboard.paste()

    req = Request(website)
    html_page = urlopen(req)

    soup = BeautifulSoup(html_page, "lxml")

    links = []
    for link in soup.findAll('a'):
        links.append(link.get('href'))
    full_link = str(links[3])
    split_link = full_link.split('/')
    print(f'The song is called {split_link[3]}')
    print(f'and the author is called {split_link[2]}')
    speak(f'The song is called {split_link[3]}')
    speak(f'and the author is called {split_link[2]}')


def close():
    pyautogui.hotkey('alt', 'f4')


def shutdown():
    speak('shutting down in 3...2..1')
    pyautogui.hotkey('ctrl', 's')
    os.system("shutdown /s /t 5")


def closepycharm():
    speak('closing pycharm')
    t.sleep(1.5)
    pyautogui.hotkey('alt', 'f4')
    t.sleep(3)
    pyautogui.leftClick(731, 426)


def dictionary():
    split_mean = text.split('does')
    index_split = split_mean[1]
    word1 = index_split.split('mean')
    final = str(word1[0])
    print(final)
    global dc
    dc = PyDictionary()
    meaning = dc.meaning(final)
    speak(meaning)
    print(meaning)


def big_dictionary():
    speak('Launching..')
    dc = PyDictionary()
    doc = docx.Document()
    doc.add_heading('English Vocabulary ', 0)

    string = input('please enter words')
    print(f'These are the words:{string} ')
    lst = string.split()
    print(lst)

    for i in lst:
        print(i)

        meaning = dc.meaning(i)
        str_meaning = str(meaning)
        print(str_meaning)
        global para
        para = doc.add_paragraph().add_run(f"{i} : {str_meaning}")
        para.font.size = Pt(18)

        doc.save(r"C:\\Users\\Dr.Wael\\Desktop\\English vocab\\english_vocab.docx")


def int_time(minute=True):
    time = today_date.strftime("%X")
    split_time = time.split(':')
    if minute is True:
        index_time = split_time[1]
    else:
        index_time = split_time[0]
    global integer_time
    if minute is True:
        integer_time = int(index_time) + 1
    else:
        integer_time = int(index_time)


def music_player(song):
    global musical
    musical = vlc.MediaPlayer(song)
    musical.audio_set_volume(25)
    musical.play()


def random_music():
    path = "C:\\\\Users\\\\Dr.Wael\\\\PycharmProjects\\\\The talking project\\\\musicfiles"
    files = os.listdir(path)
    d = random.choice(files)
    music_player(d)
    speak(f'playing {d}')
    music_play_or_stop()


def music_play_or_stop():
    x = input('type pause to pause')
    if x == 'pause':
        musical.pause()
        f = input('type play to continue')
        if f == 'play':
            musical.play()


def playmusic():
    random_music()


def click(button):  # function to click buttons
    t.sleep(2.5)
    pyautogui.press(button)


def open_minecraft():
    speak('opening minecraft')
    os.startfile(r'C:\\Users\\Dr.Wael\\AppData\\Roaming\\.tlauncher\\mcl\\Minecraft\\TL.exe')
    t.sleep(40)
    pyautogui.click(658, 444, button='left')


def recycle_bin():
    speak('recycling bin')
    winshell.recycle_bin().empty(confirm=False, show_progress=False, sound=False)


def read():  # function to read selected text
    speak('reading...')
    t.sleep(2)
    pyautogui.hotkey('ctrl', 'c')
    pyautogui.hotkey('ctrl', 'c')
    pyautogui.hotkey('ctrl', 'c')
    t.sleep(2)
    text_to_read = clipboard.paste()
    print(text_to_read)
    speak(text_to_read)


def send_message():
    int_time(minute=True)
    minute = integer_time
    int_time(minute=False)
    hour = integer_time
    print(hour, minute)

    kt.sendwhatmsg(phone_number,
                   msg,
                   hour, minute)
    t.sleep(25)
    pyautogui.click(977, 321, button='left')
    t.sleep(5)
    pyautogui.click(977, 321, button='left')
    t.sleep(2)
    click('enter')

    speak("message sent")


def call_yousef():
    speak('calling yousef')
    os.startfile(r'C:\\Users\\Dr.Wael\\AppData\\Local\\Discord\\app-1.0.9004\\discord.exe')
    t.sleep(80)
    pyautogui.click(165, 93, button='left')
    t.sleep(4)
    pyautogui.click(540, 40, button='left')  # all friends
    t.sleep(4)
    pyautogui.click(426, 202, button='left')  # yousef
    t.sleep(1)
    pyautogui.click(983, 48, button='left')  # call button
    speak('called him')


def quiz():
    speak("activating quiz mode")

    speak("mode activated sir")
    while True:
        text = get_audio(printorno=True)

        if 'break' in text:
            break
        elif "2" in text:
            button = "num2"
            try:
                speak("pressing")
                click(button)

            except ValueError:
                pass
        elif "1" in text:
            button = "num1"
            try:
                speak("pressing")
                click(button)

            except ValueError:
                pass
        elif "3" in text:
            button = "num3"
            try:
                speak("pressing")
                click(button)

            except ValueError:
                pass
        elif "4" in text:
            button = "num4"
            try:
                speak("pressing")
                click(button)

            except ValueError:
                pass


def searchmodule():
    print("searching")
    if "search" in text:
        splitted_text = text.split('search')
    elif "sir" in text:
        splitted_text = text.split('sir')
    elif "share" in text:
        splitted_text = text.split('share')
    else:
        splitted_text = ''
    print(splitted_text)
    try:
        kt.search(splitted_text[1])
        t.sleep(10)

    except ValueError:
        speak("couldn't find what to search")


def school_scheduler():
    today_date: datetime = datetime.now()
    day = today_date.strftime("%A")
    print(day)
    if day == "Sunday":
        os.startfile(r"C:\Users\Dr.Wael\PycharmProjects\The talking project\school schedule\monday.jpg")
        t.sleep(5)
    if day == "Monday":
        os.startfile(r"C:\Users\Dr.Wael\PycharmProjects\The talking project\school schedule\tuesday.jpg")
        t.sleep(5)
    if day == "Tuesday":
        os.startfile(r"C:\Users\Dr.Wael\PycharmProjects\The talking project\school schedule\wednesday.jpg")
        t.sleep(5)
    if day == "Wednesday":
        os.startfile(r"C:\Users\Dr.Wael\PycharmProjects\The talking project\school schedule\thursday.jpg")
        t.sleep(5)


stopped = False
if __name__ == '__main__':
    while True:

        while stopped is False:
            text = get_audio(printorno=True)
            if len(text) == 0:
                speak("couldn't get input")
            if "hello" in text or 'good morning' in text:
                speak("hello, how are you?")
                continue
            elif 'stop' in text:
                stopped = True

                text = get_audio(printorno=False)  # code
                speak("Waiting for listen command to continue")
                print("not listening")
                break
            elif 'close' in text:
                speak('closing')
                close()
            elif 'school' in text:
                school_scheduler()
            elif 'lyrics' in text or 'song' and 'called' in text:
                get_song_by_lyrics()
            elif 'send' in text and 'dad' in text or 'send' in text and 'that' in text:
                phone_number = "+0553880987"
                speak('what do you want to send?')
                text = get_audio(printorno=True)
                msg = text
                send_message()
            elif 'files' in text:
                file_delete()
            elif 'read' in text:
                t.sleep(2)
                read()
            elif 'destroy' in text:
                lock_keyboard()
            elif 'unlock' in text:
                unlock_keyboard()
            elif 'recycle' in text or 'empty' in text:
                recycle_bin()
            elif 'mean' in text:
                dictionary()
            elif 'shut down' in text or 'shutdown' in text:
                closepycharm()
                close()
                close()
                shutdown()
            elif 'close' in text:
                os.system("TASKKILL /F /IM MediaMonkey.exe")
                close()

            elif 'phone' in text:
                call_phone()
            elif "search" in text or "sir" in text or "share" in text:
                searchmodule()
            elif "music" in text:
                playmusic()
            elif 'dictionary' in text or 'vocab' in text:
                big_dictionary()
                t.sleep(4)
                os.startfile(r"C:\\\\Users\\\\Dr.Wael\\\\Desktop\\\\English vocab\\\\english_vocab.docx")

            elif "quiz" in text:
                quiz()
            elif 'call' in text or 'yousef' in text:
                call_yousef()

            elif 'Minecraft' in text:
                open_minecraft()
            elif "month" in text:
                month = today_date.strftime("%B")
                speak(month)
                continue
            elif "year" in text:
                year = today_date.strftime("%Y")
                speak(year)
                continue
            elif "day" in text:
                day = today_date.strftime("%A")
                speak(day)
                continue
            elif "time" in text:
                time = today_date.strftime("%X")
                speak(time)

            elif "name" in text:
                speak("hi my name is jarvis")

            else:

                i = i + 1
                print("unknown command of jarvis")
                if i == 3:
                    speak('program crashed')
                    stopped = True
                    break
                speak('Say that again,please')

        while stopped is True:
            text = get_audio(printorno=False)

            if 'listen' in text or 'hey jarvis' in text:
                speak("listening and ready")
                stopped = False
